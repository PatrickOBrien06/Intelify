from flask import Flask, render_template, request, flash, Blueprint
from flask_login import login_required
from openai import OpenAI
from website.config import API_KEY
from docx import Document
from docx.shared import Inches, Pt
from docx2pdf import convert
import json
import os

ai = Blueprint('ai', __name__, template_folder="templates")

@ai.route("/")
@ai.route("/home")
def home():
  return render_template("home.html")

@ai.route("/ai", methods=["GET", "POST"])
@login_required
def AI():
  global book_index
  if request.method == "POST":

    # Retreive prompt data for book
    topic = request.form.get("topic")
    chapters = request.form.get("chapters")
    subchapters = request.form.get("subchapters")
    pages = 1
    role = "user"

    if topic or chapters or subchapter or pages != "":

      if int(chapters) * int(subchapters) < 50:
    
        client = OpenAI(api_key=API_KEY)

        # Given the GPT API instructions on how to make the outline
        content_outline = [{"role": role, "content": f"""We are making a {pages} page book about {topic}. Create an outline for our book, which will have {chapters} chapter(s), and each chapter will have {subchapters} subchapter(s)

        Outline For Output prompt:

        'python dict with a key: chapter title, value: a single list containing each subchapter title'
        
        For example, Chapter 1 [Enter Chapter Title]: 1.1 [Enter Subchapter Title] """}]

        stream = client.chat.completions.create(
          model="gpt-3.5-turbo",
          messages=content_outline,
          temperature=0.3
        )
      
        # Grap outline of the book and make a list of each chapter
        outline = stream.choices[0].message.content

        try:
          outline = json.loads(outline)
        except:
          print("JSONDecodeError, try refreshing")
          flash("An error has occured try refreshing the page.", "danger")

        chapters = outline.keys()
        print(outline)
        
        # Open a word document to add content from the AI response
        document = Document()
        
        for chapter in chapters:
          print(chapter)
          subchapters = outline[chapter]
          document.add_heading(chapter, 1)
          for subchapter in subchapters:

            # Given the outline generate each subchapter
            subchapter_content = [{"role": role, "content": f"""We are making a subchapter for a book this subchapter is about {subchapter}."""}]

            stream = client.chat.completions.create(
              model="gpt-3.5-turbo",
              messages=subchapter_content,
              temperature=0.5,
              top_p=0.5
            )
            
            document.add_heading(subchapter, 2)
            document.add_paragraph(stream.choices[0].message.content)

        # Save document and convert into a PDF file for easy access
        try:
          docx_path = "website/static/test.docx"
          pdf_path = "website/static/book.pdf"
          input_dir = os.path.dirname(docx_path)
          output_dir = os.path.dirname(pdf_path)
          os.makedirs(input_dir, exist_ok=True)
          os.makedirs(output_dir, exist_ok=True)

          document.save(docx_path)
          convert(docx_path, pdf_path)
        
        except Exception as e:
          print(f"{str(e)} Refreshing.")
          flash("Try Refreshing the page!", "danger")

        
      else:
        print("Exceeded Page Limit!")
        flash("Exceeded Page Limit!", "danger")
    
    else: 
      print("Invalid Input")
      flash("Invalid Input!", "danger")
      

  return render_template("index.html")