from flask import Flask, render_template, request
from openai import OpenAI
from config import API_KEY
from docx import Document
from docx.shared import Inches, Pt
from docx2pdf import convert
import json


app = Flask(__name__)

@app.route("/", methods=["GET", "POST"])
def home():
  if request.method == "POST":

    # Retreive prompt data for book
    topic = request.form.get("topic")
    chapters = request.form.get("chapters")
    subchapters = request.form.get("subchapters")
    pages = request.form.get("pages")
    role = "user"
    
    client = OpenAI(api_key=API_KEY)

    # Given the GPT API instructions on how to make the outline
    content_outline = [{"role": role, "content": f"""We are making a {pages} page book about {topic}. Create an outline for our book, which will have {chapters} chapters, and each chapter will have {subchapters} subchapters

    Outline For Output prompt:

    'python dict with a key: chapter title, value: a single list containing each subchapter title'
    
    For example, Chapter 1 [Enter Chapter Title]: 1.1 [Enter Subchapter Title] """}]

    stream = client.chat.completions.create(
      model="gpt-3.5-turbo",
      messages=content_outline
    )
  
    # Grap outline of the book and make a list of each chapter
    outline = stream.choices[0].message.content
    outline = json.loads(outline)
    chapters = outline.keys()
    print(outline)
    
    # Open a word document to add content from the AI response
    document = Document()
    
    for chapter in chapters:
      print(chapter)
      subchapters = outline[chapter]
      document.add_heading(chapter, 1)
      for subchapter in subchapters:

        # Given the outline generate each subchapter
        subchapter_content = [{"role": role, "content": f"""We are making a a subchapter for a book this subchapter is about {subchapter}."""}]

        stream = client.chat.completions.create(
          model="gpt-3.5-turbo",
          messages=subchapter_content
        )
        
        document.add_heading(subchapter, 2)
        document.add_paragraph(stream.choices[0].message.content)

    # Save document and convert into a PDF file for easy access
    document.save("test.docx")
    convert("test.docx", "book.pdf")
    

  return render_template("index.html")


if __name__ == "__main__":
  app.run(debug=True)